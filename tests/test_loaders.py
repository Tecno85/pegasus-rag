from io import BytesIO

import pandas as pd
import pytest
from docx import Document
from pypdf import PdfWriter

from pegasus_rag.errors import EmptyDocumentError, ScannedDocumentError, UnsupportedFormatError
from pegasus_rag.loaders import load_document, validate_upload


def test_docx_preserves_heading_and_paragraph() -> None:
    buffer = BytesIO()
    document = Document()
    document.add_heading("Seguridad", level=1)
    document.add_paragraph("Nunca guardes una API key en el repositorio.")
    document.save(buffer)

    sections = load_document("manual.docx", buffer.getvalue())

    assert any(section.location == "Sección Seguridad, párrafo 2" for section in sections)
    assert "API key" in sections[-1].text


def test_csv_groups_rows_and_keeps_locations() -> None:
    data = b"producto,stock\nTeclado,12\nMonitor,4\n"

    sections = load_document("inventario.csv", data)

    assert len(sections) == 1
    assert sections[0].location == "filas 2-3"
    assert "producto: Teclado" in sections[0].text


def test_xlsx_keeps_sheet_name() -> None:
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        pd.DataFrame({"equipo": ["Hermes"], "estado": ["activo"]}).to_excel(
            writer, sheet_name="Squads", index=False
        )

    sections = load_document("equipos.xlsx", buffer.getvalue())

    assert sections[0].location == "Hoja Squads, filas 2-2"
    assert "Hermes" in sections[0].text


def test_rejects_unsupported_and_empty_files() -> None:
    with pytest.raises(UnsupportedFormatError):
        validate_upload("notes.txt", b"text", 10)
    with pytest.raises(EmptyDocumentError):
        validate_upload("empty.pdf", b"", 10)


def test_pdf_without_embedded_text_is_reported_as_scanned() -> None:
    buffer = BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=300, height=300)
    writer.write(buffer)

    with pytest.raises(ScannedDocumentError, match="OCR"):
        load_document("scan.pdf", buffer.getvalue())
